import pandas as pd
import os
import re
import random
import io
from flask import Blueprint, send_file
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from utils.common import normalize_damage
from utils.evaluation import classify_repair, match_priority, match_unit_price, adjust, natural_sort_key
from static.data.damage_solutions import damage_solutions
from utils.damage_ai_generator import get_damage_solution

download_bp = Blueprint('download', __name__)

@download_bp.route('/download/<table_type>')
def download_table(table_type):
    if not os.path.exists('cache.csv'):
        return "데이터가 없습니다."

    df = pd.read_csv('cache.csv')

    if table_type == 'detail':
        # 부재별 집계표 생성
        doc = Document()
        doc.add_heading('부재별 집계표', 0)

        for name, group in df.groupby('부재명'):
            doc.add_heading(f'부재명: {name}', level=1)

            # 관찰 결과 요약 추가
            unique_damages = sorted(set(group['손상내용']))
            damage_text = ', '.join(unique_damages)
            summary = f"{name}에 대한 외관조사결과에서 조사된 바와 같이, {damage_text} 등의 손상이 조사되었다."
            doc.add_paragraph(summary)

            # 부재위치 목록 가져오기
            positions = sorted(group['부재위치'].unique(), key=natural_sort_key)

            # 테이블 생성 (합계 열을 포함하여 +2)
            table = doc.add_table(rows=2, cols=3 + len(positions) * 2 + 2)
            table.style = 'Table Grid'

            # 표 스타일 설정
            for row in table.rows:
                for cell in row.cells:
                    # 셀 여백 설정
                    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                    # 텍스트 정렬 설정
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    # 폰트 크기 설정
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

            # 첫 번째 행 헤더 (부재위치)
            header_cells = table.rows[0].cells
            header_cells[0].text = '부재명'
            header_cells[1].text = '손상내용'
            header_cells[2].text = '단위'

            # 부재위치별 헤더 추가 (첫 번째 행)
            for i, pos in enumerate(positions):
                header_cells[3 + i*2].text = pos
                header_cells[4 + i*2].text = pos

            # 합계 열 헤더 추가 (첫 번째 행)
            header_cells[-2].text = '합계'
            header_cells[-1].text = '합계'

            # 두 번째 행 헤더 (손상물량/개소)
            subheader_cells = table.rows[1].cells
            subheader_cells[0].text = ''
            subheader_cells[1].text = ''
            subheader_cells[2].text = ''

            # 부재위치별 서브헤더 추가 (두 번째 행)
            for i in range(len(positions)):
                subheader_cells[3 + i*2].text = '손상물량'
                subheader_cells[4 + i*2].text = '개소'
            # 합계 열 서브헤더 추가 (두 번째 행)
            subheader_cells[-2].text = '손상물량'
            subheader_cells[-1].text = '개소'

            # 손상내용별로 데이터 그룹화
            damage_groups = group.groupby('손상내용')

            # 데이터 추가
            for damage, damage_group in damage_groups:
                row_cells = table.add_row().cells
                row_cells[0].text = str(name)
                row_cells[1].text = str(damage)
                row_cells[2].text = str(damage_group['단위'].iloc[0])  # 첫 번째 행의 단위 사용

                # 부재위치별 데이터 추가
                total_damage = 0
                total_count = 0
                for i, pos in enumerate(positions):
                    pos_data = damage_group[damage_group['부재위치'] == pos]
                    if not pos_data.empty:
                        damage_val = float(pos_data['손상물량'].iloc[0])
                        count = int(pos_data['개소'].iloc[0])
                        row_cells[3 + i*2].text = f"{damage_val:.2f}"
                        row_cells[4 + i*2].text = str(count)
                        total_damage += damage_val
                        total_count += count
                    else:
                        row_cells[3 + i*2].text = '-'
                        row_cells[4 + i*2].text = '-'

                # 합계 데이터 추가
                row_cells[-2].text = f"{total_damage:.2f}"
                row_cells[-1].text = str(total_count)

                # 새로 추가된 행의 스타일 설정
                for cell in row_cells:
                    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

            # 손상별 원인 및 대책방안 추가
            doc.add_heading('손상별 원인 및 대책방안', level=2)

            for dmg in unique_damages:
                repair_method = classify_repair(dmg)
                doc.add_paragraph(f'손상내용: {dmg} (보수방안: {repair_method})', style='Heading 3')

                try:
                    normalized_dmg = normalize_damage(dmg)

                    # 균열 관련 처리
                    if '균열' in normalized_dmg and not any(x in normalized_dmg for x in ['백태', '누수', '부']):
                        if re.search(r'균열\(?[a-zA-Z]*=?[\d.]+(mm|㎜)', normalized_dmg):
                            # 균열로 처리
                            formatted_solution = get_damage_solution("균열", name, repair_method)
                            doc.add_paragraph(formatted_solution)
                        else:
                            # 일반 균열 처리
                            formatted_solution = get_damage_solution("균열", name, repair_method)
                            doc.add_paragraph(formatted_solution)
                    else:
                        # 다른 손상 유형 처리 (damage_solutions에 없는 경우 AI로 생성)
                        formatted_solution = get_damage_solution(normalized_dmg, name, repair_method)
                        doc.add_paragraph(formatted_solution)
                except Exception:
                    pass

                doc.add_paragraph('-' * 50)  # 구분선 추가

            doc.add_page_break()  # 부재별 구분을 위한 페이지 나누기

    elif table_type == 'overall':
        # 외관조사 총괄표 생성
        overall = df.groupby(['부재명', '손상내용', '단위'])[['손상물량', '개소']].sum().reset_index()
        overall['손상물량'] = overall['손상물량'].round(2)

        doc = Document()
        doc.add_heading('외관조사 총괄표', 0)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # 표 스타일 설정
        for row in table.rows:
            for cell in row.cells:
                # 셀 여백 설정
                cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                # 텍스트 정렬 설정
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # 폰트 크기 설정
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        # 헤더 추가
        header_cells = table.rows[0].cells
        header_cells[0].text = '부재명'
        header_cells[1].text = '손상내용'
        header_cells[2].text = '단위'
        header_cells[3].text = '손상물량'
        header_cells[4].text = '개소'

        # 데이터 추가
        for _, row in overall.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['부재명'])
            row_cells[1].text = str(row['손상내용'])
            row_cells[2].text = str(row['단위'])
            row_cells[3].text = f"{float(row['손상물량']):.2f}"
            row_cells[4].text = str(row['개소'])

            # 새로 추가된 행의 스타일 설정
            for cell in row_cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    elif table_type == 'repair':
        # 보수물량표 생성 - repair 함수와 동일한 로직 사용
        doc = _generate_repair_table(df)

    elif table_type == 'cost':
        # 개략공사비표 생성 - cost 함수와 동일한 로직 사용
        doc = _generate_cost_table(df)

    # 문서를 메모리에 저장
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    # 파일 이름 설정
    if table_type == 'detail':
        filename = '부재별_집계표.docx'
    elif table_type == 'overall':
        filename = '외관조사_총괄표.docx'
    elif table_type == 'repair':
        filename = '보수물량표.docx'
    elif table_type == 'cost':
        filename = '개략공사비표.docx'
    else:
        filename = f'{table_type}_표.docx'

    return send_file(
        doc_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

def _generate_repair_table(df):
    """보수물량표 생성 함수"""
    repair = df.groupby(['부재명', '손상내용', '단위'])[['손상물량', '개소']].sum().reset_index()
    repair['보수방안'] = repair['손상내용'].apply(classify_repair)
    repair['우선순위'] = repair.apply(lambda row: match_priority(row['손상내용'], repair_method=row['보수방안']), axis=1)
    repair['단가'] = repair['손상내용'].apply(match_unit_price)
    # 보수방안이 '주의관찰'인 경우 단가를 0으로 설정
    repair.loc[repair['보수방안'] == '주의관찰', '단가'] = 0
    repair['보수물량'] = repair.apply(adjust, axis=1)
    repair['보수물량'] = repair['보수물량'].round(2)

    doc = Document()
    doc.add_heading('보수물량표', 0)

    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'

    # 표 스타일 설정
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # 헤더 추가
    header_cells = table.rows[0].cells
    header_cells[0].text = '부재명'
    header_cells[1].text = '손상내용'
    header_cells[2].text = '단위'
    header_cells[3].text = '손상물량'
    header_cells[4].text = '보수물량'
    header_cells[5].text = '개소'
    header_cells[6].text = '보수방안'
    header_cells[7].text = '우선순위'
    header_cells[8].text = '단가'

    # 데이터 추가
    for _, row in repair.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['부재명'])
        row_cells[1].text = str(row['손상내용'])
        row_cells[2].text = str(row['단위'])
        row_cells[3].text = f"{float(row['손상물량']):.2f}"
        row_cells[4].text = f"{float(row['보수물량']):.2f}"
        row_cells[5].text = str(row['개소'])
        row_cells[6].text = str(row['보수방안'])
        row_cells[7].text = str(row['우선순위'])
        row_cells[8].text = f"{int(row['단가']):,}"

        # 새로 추가된 행의 스타일 설정
        for cell in row_cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    return doc

def _generate_cost_table(df):
    """개략공사비표 생성 함수"""
    repair = df.groupby(['부재명', '손상내용', '단위'])[['손상물량', '개소']].sum().reset_index()
    repair['보수방안'] = repair['손상내용'].apply(classify_repair)
    repair['우선순위'] = repair.apply(lambda row: match_priority(row['손상내용'], repair_method=row['보수방안']), axis=1)
    repair['단가'] = repair['손상내용'].apply(match_unit_price)
    repair['보수물량'] = repair.apply(adjust, axis=1)
    repair['보수물량'] = repair['보수물량'].round(2)
    repair['개략공사비'] = repair['보수물량'] * repair['단가']

    filtered = repair[repair['보수방안'] != '주의관찰'].copy()
    result = filtered.groupby(['부재명', '보수방안', '우선순위'], dropna=False).agg({
        '보수물량': 'sum',
        '개소': 'sum',
        '손상내용': lambda x: ', '.join(sorted(set(x))),
        '단가': 'first',
        '개략공사비': 'sum'
    }).reset_index()

    doc = Document()
    doc.add_heading('개략공사비표', 0)

    # 개략공사비표
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # 표 스타일 설정
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    header_cells = table.rows[0].cells
    header_cells[0].text = '부재명'
    header_cells[1].text = '보수방안'
    header_cells[2].text = '우선순위'
    header_cells[3].text = '보수물량'
    header_cells[4].text = '개소'
    header_cells[5].text = '개략공사비'

    for _, row in result.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['부재명'])
        row_cells[1].text = str(row['보수방안'])
        row_cells[2].text = str(row['우선순위'])
        row_cells[3].text = f"{float(row['보수물량']):.2f}"
        row_cells[4].text = str(row['개소'])
        row_cells[5].text = f"{int(row['개략공사비']):,}"

        # 새로 추가된 행의 스타일 설정
        for cell in row_cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # 우선순위별 요약
    doc.add_heading('우선순위별 공사비 요약', level=1)
    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.style = 'Table Grid'

    # 표 스타일 설정
    for row in summary_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    header_cells = summary_table.rows[0].cells
    header_cells[0].text = '우선순위'
    header_cells[1].text = '순공사비'
    header_cells[2].text = '제경비(50%)'
    header_cells[3].text = '전체 공사비'

    group_by_priority = result.groupby('우선순위')['개략공사비'].sum()
    total_sum = 0

    for prio, amount in group_by_priority.items():
        soft = int(amount)
        indirect = int(soft * 0.5)
        total = soft + indirect
        total_sum += total

        row_cells = summary_table.add_row().cells
        row_cells[0].text = str(prio)
        row_cells[1].text = f"{soft:,}"
        row_cells[2].text = f"{indirect:,}"
        row_cells[3].text = f"{total:,}"

        # 새로 추가된 행의 스타일 설정
        for cell in row_cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # 총괄 개략공사비
    row_cells = summary_table.add_row().cells
    row_cells[0].text = '총괄 개략공사비'
    row_cells[1].text = ''
    row_cells[2].text = ''
    row_cells[3].text = f"{int(total_sum):,}"

    # 새로 추가된 행의 스타일 설정
    for cell in row_cells:
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

    return doc
